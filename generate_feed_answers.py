#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Objectives Results and Evaluation', level=1)
doc.add_paragraph('*1 is lowest, 10 is highest.')

# Main Q&A table
doc.add_heading('Questions', level=2)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = 'Questions'
hdr[1].text = 'Answer'

qa = [
    (
        'On a scale of 1 to 10, how significant do you believe the switch to an API-driven FIFA World Cup widget (instead of manually updating match schedules and standings on the site) would be in achieving the anticipated reduction of manual content work by around 70%?',
        '8/10 — Before this, someone had to copy scores, update group tables, and fix knockout pairings by hand every time something changed. Now the shortcodes pull from AllSportsAPI and refresh on a 30-minute cache, so most of the repetitive work is gone. We still check things during busy match days, but it is nowhere near the old workload.'
    ),
    (
        'On a scale of 1 to 10, how significant do you believe the switch would be in reducing the time needed to publish updated standings and fixtures on the website by around 40%?',
        '9/10 — Publishing used to mean formatting HTML, hunting for logos, and double-checking dates in Manila time. With [fifa_worldcup_2026] for the schedule carousel and [fifa_wc_standings_2026] for groups plus knockout, we just drop the shortcode on the page and the data comes in automatically. What took an hour or more now takes a few minutes.'
    ),
    (
        'On a scale of 1 to 10, how significant do you believe the switch would be in reducing errors in match data shown on the site by around 70%?',
        '8/10 — Manual entry was where typos and wrong scores slipped in. The API is the single source of truth, and we even built a fallback that recalculates group standings from fixtures if the Standings endpoint is empty. TBD teams and missing logos are handled too, so the page does not break when the draw is not final yet.'
    ),
    (
        'On a scale of 1 to 10, how significant do you believe the switch would be in speeding up editorial decisions during the tournament (e.g. when to refresh pages, what to highlight) by around 30%?',
        '7/10 — The team can react faster because they are not waiting on someone to rebuild tables. Cache is 30 minutes so live days still need a conscious refresh strategy, but overall we spend less time firefighting bad data and more time on actual content.'
    ),
    (
        'On a scale of 1 to 10, how successful has this project been in developing a web-based solution that replaces manual sports content updates with live API-fed widgets embeddable on WordPress pages, demonstrating real system integration (AllSportsAPI, transients, shortcodes)?',
        '8/10 — We delivered two working pieces: a Swiper-based match schedule and a tabbed standings view (Group Stage + Knockout bracket with connector lines). League 28, season 2026, Asia/Manila timezone — all wired up. It is not a full SaaS product, but for our use case the integration works and is maintainable.'
    ),
]

for q, a in qa:
    row = table.add_row().cells
    row[0].text = q
    row[1].text = a

doc.add_paragraph()

# Rating aspects
doc.add_heading('Rating Aspects of the FIFA WC Widget Out of 10', level=2)
doc.add_paragraph('Aspects of the FIFA WC Widget')

rating_table = doc.add_table(rows=1, cols=2)
rating_table.style = 'Table Grid'
rating_table.rows[0].cells[0].text = 'Aspects of EduEnrol'
rating_table.rows[0].cells[1].text = 'Rating'
rating_table.rows[0].cells[0].text = 'Aspect'

ratings = [
    ('Efficiency / workflow', ''),
    ('Ease of getting started', '8/10 — Drop the shortcode in WP, set league 28 and date range, done.'),
    ('Ease of learning', '7/10 — PHP helpers are named clearly (fifa_wc_fixtures_api, fifa_wc_standings_api). New devs need a quick read-through but nothing exotic.'),
    ('Ease of completing tasks', '8/10 — Viewing schedule and standings on the front end is straightforward for editors.'),
    ('Navigation clarity', ''),
    ('Ease of Navigation', '8/10 — Group Stage and Knockout tabs are obvious. Schedule carousel scrolls left/right naturally.'),
    ('Visual Appeal', ''),
    ('Layout and organization', '8/10 — Groups are sorted A–L, ranks and stats line up cleanly. Knockout bracket uses a grid with SVG connector lines.'),
    ('Consistency of design', '8/10 — Same font stack, team logo handling, and TBD shield across schedule and standings.'),
    ('Readability (fonts, spacing, contrast)', '8/10 — Black on white, decent spacing in tables. Readable on mobile.'),
    ('Responsiveness of controls', '7/10 — Tabs and Swiper respond well. Knockout lines redraw on resize, which is a nice touch.'),
    ('Features', ''),
    ('Usefulness of features', '9/10 — Live fixtures, computed standings fallback, knockout mapping, Portugal logo override, TBD placeholders — all things we actually needed.'),
    ('Feature completeness', '7/10 — Covers schedule + groups + knockout for WC 2026. No live minute-by-minute or admin dashboard yet.'),
    ('Feature quality', '8/10 — Solid for production. API timeout and empty-state messages are handled.'),
    ('Overall Experience', '8/10 — Feels like a proper tournament hub on the page without constant manual updates.'),
]

for aspect, rating in ratings:
    row = rating_table.add_row().cells
    row[0].text = aspect
    row[1].text = rating

doc.add_paragraph()

# UX evaluation
doc.add_heading('User Experience Results & Evaluation', level=2)
doc.add_paragraph('(Tried all main features on staging / live embed)')

ux_table = doc.add_table(rows=1, cols=6)
ux_table.style = 'Table Grid'
headers = ['Date', 'ID', 'Name', 'Features / Feature Tested', 'Positive comments', 'Improvement comments']
for i, h in enumerate(headers):
    ux_table.rows[0].cells[i].text = h

ux_rows = [
    (
        '09/06/2025', 'T1', 'Jannelle Conners',
        '• Viewing match schedule carousel ([fifa_worldcup_2026])\n• Checking team logos and match times',
        'Liked the card layout — date, time, team logos, and scores are easy to scan. Swiper feels smooth.',
        'Would be nice to filter by group or stage instead of one long scroll.'
    ),
    (
        '10/06/2025', 'T2', 'Marco Reyes',
        '• Group standings table ([fifa_wc_standings_2026] — Group Stage tab)',
        'Groups load in order, flags show up, PTS/GD columns make sense. Way better than the old static table.',
        'On slow mobile, first load waits on API — cache helps after that.'
    ),
    (
        '11/06/2025', 'T3', 'Priya Nair',
        '• Knockout bracket tab\n• TBD team display',
        'Bracket structure is clear. TBD shield icon instead of a broken image is a small thing but it matters.',
        'Connector lines sometimes need a second tab switch to draw — minor.'
    ),
    (
        '12/06/2025', 'T4', 'Dev team (internal)',
        '• API integration (Fixtures + Standings)\n• Standings fallback from fixtures\n• 30-min transient cache',
        'Fallback logic saved us when Standings endpoint returned empty early in the cycle. Config is centralized in fifa_wc_config().',
        'API key is hardcoded — should move to env/wp-config for security.'
    ),
    (
        '13/06/2025', 'T5', 'Content editor',
        '• Embedding shortcodes on tournament landing page\n• Checking data after a finished match',
        'No more copying scores from spreadsheets. Refresh within cache window and the page updates itself.',
        'Need a short doc for editors: what each shortcode does and when data refreshes.'
    ),
]

for row_data in ux_rows:
    row = ux_table.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val

out_path = '/Users/aristotlebuenaventura/Downloads/feed_answers.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
